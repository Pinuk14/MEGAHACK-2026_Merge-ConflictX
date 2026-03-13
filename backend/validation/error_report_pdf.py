from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import tempfile
import os
import uuid
from pathlib import Path

# Default logo path
DEFAULT_LOGO_PATH = r"D:\projects\ai_data_pipeline\Logo.png"
DEFAULT_WATERMARK_TEXT = "NORIXIS"


# --------------------------------------------------
# Utility: Set table cell background color
# --------------------------------------------------
def set_cell_background(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


# --------------------------------------------------
# Main Report Generator
# --------------------------------------------------
def generate_report(warning_log, error_log, output_file="validation_report", logo_path: str = None, watermark_text: str = None, watermark_alpha: float = 0.15):
    """
    Generates DOCX + PDF validation report
    Warnings and errors are grouped under each record

    Args:
        warning_log: List of warning dicts with 'record' and 'message'
        error_log: List of error dicts with 'record' and 'message'
        output_file: Base output filename (without extension)
        logo_path: path to image to place in header
        watermark_text: short watermark string to show under the logo in header
        watermark_alpha: transparency level for logo watermark (0.0-1.0, default 0.15)

    Returns:
        Tuple of (docx_path, pdf_generated_bool)
    """

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Use default logo if none provided and file exists
    if logo_path is None and Path(DEFAULT_LOGO_PATH).exists():
        logo_path = DEFAULT_LOGO_PATH
        if watermark_text is None:
            watermark_text = DEFAULT_WATERMARK_TEXT

    # Process logo with transparency for watermark effect
    watermark_temp_file = None
    if logo_path:
        try:
            from PIL import Image
            img = Image.open(logo_path)
            if img.mode != 'RGBA':
                img = img.convert('RGBA')
            
            # Apply transparency
            alpha = img.split()[3]
            alpha = alpha.point(lambda p: int(p * watermark_alpha))
            img.putalpha(alpha)
            
            # Save to temp file
            watermark_temp_file = f"temp_watermark_{uuid.uuid4().hex}.png"
            img.save(watermark_temp_file, 'PNG')
            logo_path = watermark_temp_file
        except ImportError:
            print("⚠️  PIL not available, using logo without transparency adjustment")
        except Exception as e:
            print(f"⚠️  Error processing logo: {e}")

    # Header watermark (logo + optional watermark text)
    if logo_path or watermark_text:
        for section in doc.sections:
            header = section.header
            # clear existing header content
            for p in list(header.paragraphs):
                p.clear()

            hdr_para = header.add_paragraph()
            hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                if logo_path:
                    run_logo = hdr_para.add_run()
                    run_logo.add_picture(logo_path, width=Inches(5.5))
            except Exception as e:
                print(f"⚠️ Logo not added to header (logo_path={logo_path}): {e}")

            if watermark_text:
                wt_para = header.add_paragraph()
                wt_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                wt_run = wt_para.add_run(watermark_text)
                wt_run.font.size = Pt(14)
                wt_run.font.color.rgb = RGBColor(180, 180, 180)
                wt_run.font.italic = True

    # --------------------------------------------------
    # Title
    # --------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Validation Report")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 84, 144)

    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generated on {timestamp}").font.size = Pt(10)

    doc.add_paragraph()

    # --------------------------------------------------
    # Summary
    # --------------------------------------------------
    total_warning_records = len(set(w["record_index"] for w in warning_log))
    total_error_records = len(set(e["record_index"] for e in error_log))
    total_errors = sum(len(e["errors"]) for e in error_log)

    summary = doc.add_table(rows=2, cols=3)
    summary.style = "Light Grid Accent 1"
    headers = ["Records with Warnings", "Records with Errors", "Total Errors"]
    values = [total_warning_records, total_error_records, total_errors]

    for i, cell in enumerate(summary.rows[0].cells):
        cell.text = headers[i]
        set_cell_background(cell, "2c5aa0")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, cell in enumerate(summary.rows[1].cells):
        cell.text = str(values[i])
        set_cell_background(cell, "f0f4f8")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(26, 84, 144)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # --------------------------------------------------
    # Merge warnings + errors per record
    # --------------------------------------------------
    record_map = {}

    for w in warning_log:
        record_map.setdefault(w["record_index"], {"warnings": [], "errors": []})
        record_map[w["record_index"]]["warnings"].append(w["warning"])

    for e in error_log:
        record_map.setdefault(e["record_index"], {"warnings": [], "errors": []})
        record_map[e["record_index"]]["errors"].extend(e["errors"])

    # --------------------------------------------------
    # Per Record Sections
    # --------------------------------------------------
    for record_index in sorted(record_map.keys()):
        data = record_map[record_index]

        # Record header
        header = doc.add_table(rows=1, cols=1)
        cell = header.rows[0].cells[0]
        cell.text = f"Record {record_index}"
        set_cell_background(cell, "2c5aa0")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(14)
                r.font.color.rgb = RGBColor(255, 255, 255)

        # ---------------- WARNINGS ----------------
        if data["warnings"]:
            warn_title = doc.add_paragraph("⚠ Warnings")
            warn_title.runs[0].font.size = Pt(14)
            warn_title.runs[0].font.bold = True
            warn_title.runs[0].font.color.rgb = RGBColor(176, 58, 46)

            warn_table = doc.add_table(rows=len(data["warnings"]), cols=2)
            for i, warn in enumerate(data["warnings"]):
                row = warn_table.rows[i].cells
                row[0].text = "!"
                set_cell_background(row[0], "fff3cd")  # Yellow
                row[1].text = warn
                set_cell_background(row[1], "fff3cd")

        # ---------------- ERRORS ----------------
        if data["errors"]:
            err_title = doc.add_paragraph("❌ Errors")
            err_title.runs[0].font.size = Pt(14)
            err_title.runs[0].font.bold = True
            err_title.runs[0].font.color.rgb = RGBColor(192, 57, 43)

            err_table = doc.add_table(rows=len(data["errors"]), cols=2)
            for i, err in enumerate(data["errors"]):
                row = err_table.rows[i].cells
                row[0].text = f"{i+1}."
                set_cell_background(row[0], "fdecea")  # Red
                row[1].text = err
                set_cell_background(row[1], "fdecea")

        doc.add_paragraph()

    # --------------------------------------------------
    # Footer
    # --------------------------------------------------
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("End of Report").font.size = Pt(10)

    # --------------------------------------------------
    # Save + Convert
    # --------------------------------------------------
    docx_file = f"{output_file}.docx"
    doc.save(docx_file)
    print(f"✅ DOCX generated: {docx_file}")

    # Always try to convert to PDF
    pdf_success = False
    try:
        from docx2pdf import convert
        pdf_file = f"{output_file}.pdf"
        convert(docx_file, pdf_file)
        print(f"✅ PDF generated: {pdf_file}")
        pdf_success = True
    except ImportError:
        print("⚠️  PDF conversion skipped: docx2pdf not installed")
        print("    Install with: pip install docx2pdf")
    except Exception as e:
        print(f"⚠️  PDF conversion failed: {str(e)}")
        print("    DOCX file is still available")
    
    # Cleanup temp watermark file
    if watermark_temp_file and os.path.exists(watermark_temp_file):
        try:
            os.remove(watermark_temp_file)
            print(f"🗑️  Cleaned up temp watermark file")
        except Exception as e:
            print(f"⚠️  Failed to remove temp file: {e}")
    
    return docx_file, pdf_success


# --------------------------------------------------
# Example Run
# --------------------------------------------------
if __name__ == "__main__":
    warning_log = [
        {"record_index": 1, "warning": "page_count was null → auto-set to 0"},
        {"record_index": 3, "warning": "Published_date missing → auto-filled with 2025-01-03"}
    ]

    error_log = [
        {
            "record_index": 1,
            "errors": ["usability out of allowed range"]
        },
        {
            "record_index": 2,
            "errors": [
                "Title must be a non-empty string",
                "accuracy out of allowed range"
            ]
        }
    ]

    generate_report(warning_log, error_log)
